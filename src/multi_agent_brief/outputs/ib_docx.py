"""Markdown-to-DOCX converter with professional styling.

Converts Markdown documents into styled Word documents with:
- Heading hierarchy (H1-H4) with professional color scheme
- Tables (pipe and tab-separated), with wide-table fallback
- Ordered/unordered lists with nesting
- Inline formatting (bold, italic, strikethrough, code, links)
- Blockquotes, code blocks, horizontal rules
- Automatic cover page and page-numbered footer

Based on the standalone ``scripts/md_to_ib_docx.py`` script.
"""

from __future__ import annotations

import re
import warnings
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color scheme ────────────────────────────────────────────────
COLORS = {
    "primary": "003A70",   # deep blue — cover / table header / emphasis
    "dark": "002147",      # dark navy — section headings / footer / bold
    "neutral": "58595B",   # medium gray — body text
    "light_bg": "F2F2F2",  # ice gray — alternating table rows / code bg
    "negative": "9B2743",  # burgundy — inline code / risk highlights
    "positive": "2D6A4F",  # forest green — positive signals (reserved)
}

LATIN_FONT = "Arial"
EAST_ASIA_FONT = "宋体"  # 宋体

DEFAULT_FOOTER = "Generated Brief"


# ── XML helpers ─────────────────────────────────────────────────

def _hex_to_rgb(hex_str: str) -> RGBColor:
    return RGBColor(*bytes.fromhex(hex_str))


def _set_cell_shading(cell, hex_color: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_paragraph_shading(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_paragraph_left_border(paragraph, hex_color: str, size_eighths: int = 16):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_paragraph_bottom_border(paragraph, hex_color: str, size_eighths: int = 6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_run_eastasia_font(run, font_name: str):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_eastasia_font(style, font_name: str):
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:eastAsia"), font_name)


# ── Document styles ─────────────────────────────────────────────

def _setup_document_styles(doc, font_name: str):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _hex_to_rgb(COLORS["neutral"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3
    _set_style_eastasia_font(normal, font_name)

    heading_specs = [
        (1, 17, "dark", 14, 8),
        (2, 14, "primary", 12, 6),
        (3, 12, "primary", 10, 4),
        (4, 11, "dark", 8, 4),
    ]
    for level, size, color_key, before, after in heading_specs:
        h = doc.styles[f"Heading {level}"]
        h.font.name = LATIN_FONT
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = _hex_to_rgb(COLORS[color_key])
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.keep_with_next = True
        _set_style_eastasia_font(h, font_name)


# ── Markdown parser ─────────────────────────────────────────────

_RE_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
_RE_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
_RE_FENCE = re.compile(r"^(\s*)(```|~~~)\s*(\S*)\s*$")
_RE_HRULE = re.compile(r"^\s*([-*_])(\s*\1){2,}\s*$")
_RE_UL = re.compile(r"^(\s*)([-*+])\s+(.+)$")
_RE_OL = re.compile(r"^(\s*)(\d+)[.)]\s+(.+)$")
_RE_QUOTE = re.compile(r"^\s*>\s?(.*)$")

# Chinese-style section headings (common in LLM Chinese reports)
_RE_CN_H2 = re.compile(r"^[一二三四五六七八九十百千]+[、.．]\s*(.+)$")
_RE_CN_H3 = re.compile(r"^(\d+\.\d+)\s+(.+)$")
_RE_CN_H4 = re.compile(r"^（[一二三四五六七八九十]+）\s*(.+)$")


def _is_tab_row(line: str) -> bool:
    return line.count("\t") >= 2 and "|" not in line and bool(line.strip())


def _tab_rows_to_pipe(rows: list) -> tuple:
    split_rows = [[c.strip() for c in r.split("\t")] for r in rows]
    n = max(len(r) for r in split_rows)

    def to_pipe(cells):
        while len(cells) < n:
            cells.append("")
        return "| " + " | ".join(cells[:n]) + " |"

    return [to_pipe(r) for r in split_rows], "| " + " | ".join(["---"] * n) + " |"


def _is_implicit_heading(line: str, next_line: str) -> bool:
    s = line.strip()
    if not s or len(s) > 15:
        return False
    if any(p in s for p in "，。、：；！？()（）【】—…·"):
        return False
    return bool(next_line and (_RE_UL.match(next_line) or _RE_OL.match(next_line)))


def _is_any_heading(line: str, next_line: str) -> bool:
    return bool(
        _RE_HEADING.match(line)
        or _RE_CN_H2.match(line)
        or _RE_CN_H3.match(line)
        or _RE_CN_H4.match(line)
        or _is_implicit_heading(line, next_line)
    )


def parse_markdown(md_text: str) -> list:
    """Parse Markdown text into a list of blocks.

    Supports: headings, tab/pipe tables, fenced code, blockquotes,
    ordered/unordered lists (with nesting), horizontal rules,
    Chinese-style section headings.
    """
    lines = md_text.split("\n")
    blocks: list = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        next_line = lines[i + 1] if i + 1 < n else ""

        # blank line
        if not line.strip():
            i += 1
            continue

        # Markdown # heading
        m = _RE_HEADING.match(line)
        if m:
            blocks.append(("heading", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        # Fenced code block
        m = _RE_FENCE.match(line)
        if m:
            fence = m.group(2)
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith(fence):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(code_lines)))
            continue

        # Horizontal rule
        if _RE_HRULE.match(line):
            blocks.append(("hr",))
            i += 1
            continue

        # Tab-separated table
        if _is_tab_row(line):
            tab_rows = [line]
            i += 1
            while i < n and _is_tab_row(lines[i]):
                tab_rows.append(lines[i])
                i += 1
            if len(tab_rows) >= 2:
                pipe_rows, sep = _tab_rows_to_pipe(tab_rows)
                blocks.append(("table", pipe_rows, sep))
            else:
                blocks.append(("paragraph", tab_rows[0].replace("\t", "  ")))
            continue

        # Pipe table
        if "|" in line and i + 1 < n and _RE_TABLE_SEP.match(lines[i + 1]):
            table_lines = [line]
            sep_line = lines[i + 1]
            i += 2
            while i < n and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines, sep_line))
            continue

        # Blockquote
        if _RE_QUOTE.match(line):
            quote_lines = []
            while i < n and _RE_QUOTE.match(lines[i]):
                quote_lines.append(_RE_QUOTE.match(lines[i]).group(1))
                i += 1
            blocks.append(("quote", "\n".join(quote_lines).strip()))
            continue

        # Chinese-style headings
        m = _RE_CN_H2.match(line)
        if m:
            blocks.append(("heading", 1, line.strip()))
            i += 1
            continue

        m = _RE_CN_H3.match(line)
        if m:
            blocks.append(("heading", 3, line.strip()))
            i += 1
            continue

        m = _RE_CN_H4.match(line)
        if m:
            blocks.append(("heading", 4, line.strip()))
            i += 1
            continue

        # Table/figure caption
        if re.match(r"^[表图]\s*\d+[：:]\s*.+$", line.strip()):
            next_real = next((lines[j] for j in range(i + 1, n) if lines[j].strip()), "")
            if _is_tab_row(next_real) or ("|" in next_real):
                blocks.append(("caption", line.strip()))
                i += 1
                continue

        # Implicit heading (short line followed by list)
        if _is_implicit_heading(line, next_line):
            blocks.append(("heading", 3, line.strip()))
            i += 1
            continue

        # Triangle bullet paragraph
        if line.strip().startswith("▸"):
            blocks.append(("paragraph", line.strip()))
            i += 1
            continue

        # List (ordered or unordered)
        if _RE_UL.match(line) or _RE_OL.match(line):
            items = []
            while i < n:
                m_ul = _RE_UL.match(lines[i])
                m_ol = _RE_OL.match(lines[i])
                if m_ul:
                    items.append(("ul", len(m_ul.group(1).expandtabs(4)), m_ul.group(3)))
                    i += 1
                elif m_ol:
                    items.append(("ol", len(m_ol.group(1).expandtabs(4)), m_ol.group(3), int(m_ol.group(2))))
                    i += 1
                elif lines[i].startswith((" ", "\t")) and lines[i].strip() and items:
                    prev = items[-1]
                    items[-1] = (prev[0], prev[1], prev[2] + " " + lines[i].strip())
                    i += 1
                elif not lines[i].strip():
                    nxt = lines[i + 1] if i + 1 < n else ""
                    if _RE_UL.match(nxt) or _RE_OL.match(nxt):
                        i += 1
                    else:
                        break
                else:
                    break
            blocks.append(("list", items))
            continue

        # Plain paragraph (accumulate until next block boundary)
        para_lines = [line]
        i += 1
        while i < n:
            l = lines[i]
            nxt = lines[i + 1] if i + 1 < n else ""
            if not l.strip():
                break
            if _RE_FENCE.match(l) or _RE_HRULE.match(l):
                break
            if _RE_UL.match(l) or _RE_OL.match(l) or _RE_QUOTE.match(l):
                break
            if "|" in l and i + 1 < n and _RE_TABLE_SEP.match(lines[i + 1]):
                break
            if _is_tab_row(l):
                break
            if _is_any_heading(l, nxt):
                break
            if re.match(r"^[表图]\s*\d+[：:]\s*.+$", l.strip()):
                break
            para_lines.append(l)
            i += 1
        blocks.append(("paragraph", " ".join(p.strip() for p in para_lines)))

    return blocks


# ── Inline formatting ───────────────────────────────────────────

_INLINE_PATTERN = re.compile(
    r"(\*\*[^*\n]+?\*\*"        # **bold**
    r"|~~[^~\n]+?~~"            # ~~strike~~
    r"|`[^`\n]+?`"              # `code`
    r"|\[[^\]]+\]\([^)]+\)"    # [text](url)
    r"|\*[^*\n]+?\*)"           # *italic*
)


def _add_inline(paragraph, text: str, font_name: str, base_color: str | None = None):
    if not text:
        return
    parts = _INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = _hex_to_rgb(COLORS["dark"])
        elif part.startswith("~~") and part.endswith("~~") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.font.strike = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = LATIN_FONT
            _set_run_eastasia_font(run, font_name)
            run.font.size = Pt(9.5)
            run.font.color.rgb = _hex_to_rgb(COLORS["negative"])
        elif part.startswith("[") and part.endswith(")") and "](" in part:
            m_link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m_link:
                run = paragraph.add_run(m_link.group(1))
                run.font.color.rgb = _hex_to_rgb(COLORS["primary"])
                run.font.underline = True
            else:
                run = paragraph.add_run(part)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if base_color:
                run.font.color.rgb = _hex_to_rgb(base_color)

        _set_run_eastasia_font(run, font_name)


# ── Block renderers ─────────────────────────────────────────────

def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _parse_table_alignment(sep_line: str) -> list:
    aligns = []
    for piece in _parse_table_row(sep_line):
        piece = piece.strip()
        if piece.startswith(":") and piece.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif piece.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            aligns.append(WD_ALIGN_PARAGRAPH.LEFT)
    return aligns


def _render_table(doc, table_lines, sep_line, font_name: str):
    headers = _parse_table_row(table_lines[0])
    if not headers:
        return
    if len(headers) > 4:
        _render_wide_table_as_sections(doc, table_lines, sep_line, font_name)
        return
    n_cols = len(headers)

    aligns = _parse_table_alignment(sep_line)
    while len(aligns) < n_cols:
        aligns.append(WD_ALIGN_PARAGRAPH.LEFT)

    rows = []
    for line in table_lines[1:]:
        row = _parse_table_row(line)
        if len(row) < n_cols:
            row = row + [""] * (n_cols - len(row))
        else:
            row = row[:n_cols]
        rows.append(row)

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = LATIN_FONT
        _set_run_eastasia_font(run, font_name)
        _set_cell_shading(cell, COLORS["primary"])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = aligns[j]
            _add_inline(p, row[j], font_name)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i % 2 == 1:
                _set_cell_shading(cell, COLORS["light_bg"])


def _render_wide_table_as_sections(doc, table_lines, sep_line, font_name: str):
    """Render wide tables (5+ columns) as compact row sections."""
    headers = _parse_table_row(table_lines[0])
    rows = []
    for line in table_lines[1:]:
        row = _parse_table_row(line)
        if len(row) < len(headers):
            row = row + [""] * (len(headers) - len(row))
        rows.append(row[: len(headers)])

    for idx, row in enumerate(rows):
        title = row[0].strip() if row else f"Item {idx + 1}"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6 if idx == 0 else 8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = _hex_to_rgb(COLORS["dark"])
        run.font.name = LATIN_FONT
        _set_run_eastasia_font(run, font_name)

        compact_rows = []
        for header, value in zip(headers[1:], row[1:]):
            value = value.strip()
            if value:
                compact_rows.append([header.strip(), value])
        if not compact_rows:
            continue
        table = doc.add_table(rows=len(compact_rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for r_idx, (header, value) in enumerate(compact_rows):
            h_cell = table.rows[r_idx].cells[0]
            v_cell = table.rows[r_idx].cells[1]
            h_cell.text = ""
            v_cell.text = ""
            hp = h_cell.paragraphs[0]
            vp = v_cell.paragraphs[0]
            hr = hp.add_run(header)
            hr.bold = True
            hr.font.size = Pt(9.5)
            hr.font.color.rgb = _hex_to_rgb(COLORS["primary"])
            hr.font.name = LATIN_FONT
            _set_run_eastasia_font(hr, font_name)
            _add_inline(vp, value, font_name)
            for run in vp.runs:
                if run.font.size is None:
                    run.font.size = Pt(9.5)
            h_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            v_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                _set_cell_shading(h_cell, COLORS["light_bg"])
                _set_cell_shading(v_cell, COLORS["light_bg"])


_BULLETS = ["•", "○", "▪", "·", "–"]


def _render_list(doc, items, font_name: str):
    if not items:
        return
    min_indent = min(it[1] for it in items)
    counters: dict = {}
    last_kind_at: dict = {}

    for item in items:
        kind, indent, content = item[:3]
        source_number = item[3] if len(item) > 3 else None
        level = max(0, min((indent - min_indent) // 2, 4))

        for lvl in [k for k in counters if k > level]:
            counters.pop(lvl, None)
            last_kind_at.pop(lvl, None)

        if last_kind_at.get(level) != kind:
            counters[level] = source_number if kind == "ol" and source_number is not None else 1
            last_kind_at[level] = kind

        if kind == "ol":
            number = source_number if source_number is not None else counters[level]
            label = f"{number}.  "
            counters[level] = number + 1
        else:
            label = _BULLETS[level % len(_BULLETS)] + "  "

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3

        bullet_run = p.add_run(label)
        bullet_run.font.name = LATIN_FONT
        bullet_run.font.size = Pt(10.5)
        if kind == "ol":
            bullet_run.font.color.rgb = _hex_to_rgb(COLORS["primary"])
            bullet_run.bold = True
        _set_run_eastasia_font(bullet_run, font_name)

        _add_inline(p, content, font_name)


def _render_caption(doc, text: str, font_name: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = _hex_to_rgb(COLORS["neutral"])
    run.font.name = LATIN_FONT
    _set_run_eastasia_font(run, font_name)


def _render_code_block(doc, code: str, font_name: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = 1.15
    _set_paragraph_shading(p, COLORS["light_bg"])
    run = p.add_run(code)
    run.font.name = LATIN_FONT
    _set_run_eastasia_font(run, font_name)
    run.font.size = Pt(9)
    run.font.color.rgb = _hex_to_rgb(COLORS["dark"])


def _render_quote(doc, text: str, font_name: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _set_paragraph_left_border(p, COLORS["primary"])
    _add_inline(p, text, font_name, base_color=COLORS["neutral"])


def _render_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _set_paragraph_bottom_border(p, COLORS["primary"])


# ── Cover / footer ──────────────────────────────────────────────

def _add_cover(doc, title: str, subtitle: str | None, font_name: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(COLORS["primary"])
    run.font.name = LATIN_FONT
    _set_run_eastasia_font(run, font_name)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(12)
        sr = p2.add_run(subtitle)
        sr.font.size = Pt(11)
        sr.font.color.rgb = _hex_to_rgb(COLORS["neutral"])
        sr.font.italic = True
        sr.font.name = LATIN_FONT
        _set_run_eastasia_font(sr, font_name)

    _render_hr(doc)


def _add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def _add_footer(doc, footer_text: str | None, font_name: str):
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = footer_text or DEFAULT_FOOTER
    run = fp.add_run(text + " | ")
    run.font.size = Pt(8)
    run.font.color.rgb = _hex_to_rgb(COLORS["dark"])
    run.font.name = LATIN_FONT
    _set_run_eastasia_font(run, font_name)
    _add_page_field(fp)


# ── Quality check ───────────────────────────────────────────────

def _check_input_quality(md_text: str):
    if "�" in md_text:
        n = md_text.count("�")
        warnings.warn(
            f"Input Markdown contains {n} U+FFFD replacement characters. "
            f"Source text was corrupted upstream — this tool preserves them as-is.",
            stacklevel=2,
        )


# ── Public API ──────────────────────────────────────────────────

def convert(
    md_path: str | Path,
    docx_path: str | Path,
    *,
    title: str | None = None,
    subtitle: str | None = None,
    footer: str | None = None,
    font: str | None = None,
) -> Path:
    """Convert a Markdown file to a styled DOCX document.

    Args:
        md_path: Path to the input Markdown file.
        docx_path: Path for the output .docx file.
        title: Document title. If *None*, extracted from the first ``#`` heading.
        subtitle: Optional subtitle/date line below the title.
        footer: Footer text. Defaults to ``"Generated Brief"``.
        font: East-Asian font name. Defaults to ``宋体``.

    Returns:
        The *docx_path* as a :class:`~pathlib.Path`.
    """
    md_text = Path(md_path).read_text(encoding="utf-8")
    _check_input_quality(md_text)

    if title is None:
        m = re.search(r"^#\s+(.+)$", md_text, re.MULTILINE)
        if m:
            title = m.group(1).strip()
        else:
            for line in md_text.splitlines():
                line = line.strip()
                if line and not line.startswith("---"):
                    title = line
                    break
            else:
                title = Path(md_path).stem

    if font is None:
        font = EAST_ASIA_FONT

    doc = Document()
    _setup_document_styles(doc, font)
    _add_cover(doc, title, subtitle, font)

    # Remove the title line from body to avoid duplication with cover
    body_lines = md_text.splitlines()
    for idx, line in enumerate(body_lines):
        if not line.strip():
            continue
        if line.strip().lstrip("#").strip() == title:
            body_lines.pop(idx)
        break

    blocks = parse_markdown("\n".join(body_lines))
    skipped_title_block = False
    for block in blocks:
        btype = block[0]
        if btype == "heading":
            level, text = block[1], block[2]
            if not skipped_title_block and text == title:
                skipped_title_block = True
                continue
            if re.match(r"^[一二三四五六七八九十]+[、.．]", text.strip()):
                level = 1
            doc.add_heading(text, level=min(level, 4))
        elif btype == "paragraph":
            if not skipped_title_block and block[1].strip() == title:
                skipped_title_block = True
                continue
            p = doc.add_paragraph()
            _add_inline(p, block[1], font)
        elif btype == "list":
            _render_list(doc, block[1], font)
        elif btype == "table":
            _render_table(doc, block[1], block[2], font)
        elif btype == "code":
            _render_code_block(doc, block[1], font)
        elif btype == "caption":
            _render_caption(doc, block[1], font)
        elif btype == "quote":
            _render_quote(doc, block[1], font)
        elif btype == "hr":
            _render_hr(doc)

    _add_footer(doc, footer, font)

    out = Path(docx_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
