from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

from multi_agent_brief.tools.draft_cleanup import strip_claim_citations
from multi_agent_brief.outputs.naming import render_output_stem
from multi_agent_brief.outputs.source_appendix import (
    SourceAppendixResult,
    build_source_appendix,
)

_SRC_MARKER_RE = re.compile(r"\[src:[^\]]*\]")


@dataclass
class FinalizeResult:
    """Result of the reader-facing delivery finalization step."""

    status: str
    audited_brief: str
    reader_brief: str
    named_reader_brief: str = ""
    reader_docx: str = ""
    named_reader_docx: str = ""
    docx_generation: str = "not_requested"
    stripped_src_marker_count: int = 0
    source_appendix: str = ""
    source_appendix_generation: str = "not_requested"
    source_appendix_requested_by: str = "none"
    source_appendix_mode: str = "separate"
    source_appendix_source_count: int = 0
    source_appendix_cited_claim_count: int = 0
    source_appendix_resolved_claim_count: int = 0
    source_appendix_warnings: list[str] | None = None

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        if data["source_appendix_warnings"] is None:
            data["source_appendix_warnings"] = []
        return data


def finalize_reader_outputs(
    *,
    output_dir: str | Path,
    project_name: str,
    output_formats: list[str] | tuple[str, ...] | None = None,
    output_footer: str = "",
    output_named_outputs: bool = True,
    output_filename_template: str = "",
    output_filename_tokens: dict[str, str] | None = None,
    docx_template: str = "default",
    source_appendix_config: dict[str, Any] | None = None,
) -> FinalizeResult:
    """Regenerate reader-facing artifacts from internal audited markdown.

    Agent-assisted workflows write or rewrite ``output/intermediate/audited_brief.md``
    before reader-facing delivery artifacts are rendered.
    This function is the final delivery gate: it preserves the cited audited
    artifact for auditability, then writes reader-facing Markdown/DOCX outputs as
    deterministic ``strip_claim_citations(audited_brief)`` derivatives.
    """
    out = Path(output_dir)
    intermediate_dir = out / "intermediate"
    audited_path = intermediate_dir / "audited_brief.md"
    if not audited_path.exists():
        raise FileNotFoundError(
            f"Audited brief not found: {audited_path}. "
            "Run prepare/audit first or write output/intermediate/audited_brief.md."
        )

    out.mkdir(parents=True, exist_ok=True)
    intermediate_dir.mkdir(parents=True, exist_ok=True)

    audited_markdown = audited_path.read_text(encoding="utf-8")
    stripped_count = len(_SRC_MARKER_RE.findall(audited_markdown))
    base_reader_markdown = strip_claim_citations(audited_markdown)
    formats = set(output_formats or ["markdown"])
    appendix_request = _source_appendix_request(
        output_formats=formats,
        source_appendix_config=source_appendix_config or {},
    )
    appendix_path = out / "source_appendix.md"
    if appendix_path.exists():
        appendix_path.unlink()
    appendix_result = _maybe_generate_source_appendix(
        audited_markdown=audited_markdown,
        ledger_path=intermediate_dir / "claim_ledger.json",
        appendix_path=appendix_path,
        requested_by=appendix_request["requested_by"],
        explicit=bool(appendix_request["explicit"]),
    )
    reader_markdown = base_reader_markdown
    if appendix_request["mode"] == "append" and appendix_result.markdown and appendix_result.source_count:
        reader_markdown = base_reader_markdown.rstrip() + "\n\n" + appendix_result.markdown

    brief_path = out / "brief.md"
    brief_path.write_text(reader_markdown, encoding="utf-8")

    named_brief_path: Path | None = None
    if output_named_outputs:
        tokens = dict(output_filename_tokens or {})
        tokens.setdefault("project_name", project_name)
        tokens.setdefault("title", project_name)
        named_stem = render_output_stem(output_filename_template, tokens) if output_filename_template else ""
        if named_stem:
            named_brief_path = out / f"{named_stem}.md"
            if named_brief_path != brief_path:
                named_brief_path.write_text(reader_markdown, encoding="utf-8")

    docx_status = "not_requested"
    docx_path = out / "brief.docx"
    named_docx_path: Path | None = None
    if "docx" in formats:
        # Avoid leaving a stale rendered file that may still contain internal
        # [src:CLAIM_ID] markers when regeneration fails or dependencies are missing.
        if docx_path.exists():
            docx_path.unlink()
        if named_brief_path is not None:
            possible_named_docx = named_brief_path.with_suffix(".docx")
            if possible_named_docx.exists():
                possible_named_docx.unlink()
        try:
            from multi_agent_brief.outputs.ib_docx import convert

            convert(
                brief_path,
                docx_path,
                title=project_name,
                footer=output_footer or None,
                template=docx_template or "default",
            )
            docx_status = "generated"
            if named_brief_path is not None and named_brief_path.stem != "brief":
                named_docx_path = named_brief_path.with_suffix(".docx")
                shutil.copyfile(docx_path, named_docx_path)
        except ImportError:
            docx_status = "skipped_missing_dependency"
        except Exception:
            docx_status = "failed"
            raise

    result = FinalizeResult(
        status="pass",
        audited_brief=str(audited_path),
        reader_brief=str(brief_path),
        named_reader_brief=str(named_brief_path or ""),
        reader_docx=str(docx_path) if docx_path.exists() else "",
        named_reader_docx=str(named_docx_path or ""),
        docx_generation=docx_status,
        stripped_src_marker_count=stripped_count,
        source_appendix=str(appendix_path) if appendix_result.markdown and appendix_path.exists() else "",
        source_appendix_generation=appendix_result.status,
        source_appendix_requested_by=str(appendix_request["requested_by"]),
        source_appendix_mode=str(appendix_request["mode"]),
        source_appendix_source_count=appendix_result.source_count,
        source_appendix_cited_claim_count=appendix_result.cited_claim_count,
        source_appendix_resolved_claim_count=appendix_result.resolved_claim_count,
        source_appendix_warnings=appendix_result.warnings,
    )

    _assert_reader_artifact_clean(brief_path)
    if named_brief_path and named_brief_path.exists():
        _assert_reader_artifact_clean(named_brief_path)
    if appendix_path.exists():
        _assert_reader_artifact_clean(appendix_path)
    if docx_path.exists():
        _assert_docx_artifact_clean(docx_path)
    if named_docx_path and named_docx_path.exists():
        _assert_docx_artifact_clean(named_docx_path)

    report_path = intermediate_dir / "finalize_report.json"
    report_path.write_text(json.dumps(result.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    _update_audit_report_metadata(
        intermediate_dir / "audit_report.json",
        result,
        named_brief_path=named_brief_path,
    )
    return result


def _assert_reader_artifact_clean(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    if _SRC_MARKER_RE.search(text):
        raise RuntimeError(f"Reader-facing artifact still contains [src:...] marker: {path}")


def _assert_docx_artifact_clean(path: Path) -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return
    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    if _SRC_MARKER_RE.search(text + "\n" + table_text):
        raise RuntimeError(f"Reader-facing DOCX still contains [src:...] marker: {path}")


def _update_audit_report_metadata(
    audit_report_path: Path,
    result: FinalizeResult,
    *,
    named_brief_path: Path | None,
) -> None:
    if not audit_report_path.exists():
        return
    try:
        payload = json.loads(audit_report_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return
    metadata = payload.setdefault("metadata", {})
    metadata["reader_brief_artifact"] = result.reader_brief
    metadata["reader_brief_transform"] = "strip_claim_citations"
    metadata["reader_brief_finalized"] = True
    metadata["reader_brief_stripped_src_marker_count"] = result.stripped_src_marker_count
    metadata["finalize_report_artifact"] = str(audit_report_path.parent / "finalize_report.json")
    metadata["docx_generation"] = result.docx_generation
    metadata["source_appendix_generation"] = result.source_appendix_generation
    if result.source_appendix:
        metadata["source_appendix_artifact"] = result.source_appendix
    if result.reader_docx:
        metadata["rendered_docx_path"] = result.reader_docx
    if named_brief_path:
        metadata["named_reader_brief_artifact"] = str(named_brief_path)
    audit_report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _source_appendix_request(
    *,
    output_formats: set[str],
    source_appendix_config: dict[str, Any],
) -> dict[str, Any]:
    config_enabled = _as_bool(source_appendix_config.get("enabled"), False)
    if config_enabled:
        requested_by = "config"
        explicit = True
    elif "source_appendix" in output_formats:
        requested_by = "source_appendix"
        explicit = True
    elif "source_map" in output_formats:
        requested_by = "legacy_source_map"
        explicit = False
    else:
        requested_by = "none"
        explicit = False
    mode = str(source_appendix_config.get("mode") or "separate").strip().lower()
    if mode not in {"separate", "append"}:
        mode = "separate"
    return {
        "requested_by": requested_by,
        "explicit": explicit,
        "mode": mode,
    }


def _maybe_generate_source_appendix(
    *,
    audited_markdown: str,
    ledger_path: Path,
    appendix_path: Path,
    requested_by: str,
    explicit: bool,
) -> SourceAppendixResult:
    if requested_by == "none":
        return SourceAppendixResult(status="not_requested")
    if not ledger_path.exists():
        if explicit:
            raise FileNotFoundError(
                f"Claim Ledger not found for explicit source appendix request: {ledger_path}"
            )
        return SourceAppendixResult(
            status="skipped_missing_ledger",
            warnings=["Source appendix skipped because claim_ledger.json was missing."],
        )
    try:
        result = build_source_appendix(
            audited_markdown=audited_markdown,
            ledger_path=ledger_path,
        )
    except (json.JSONDecodeError, TypeError, ValueError) as exc:
        if explicit:
            raise ValueError(f"Claim Ledger is malformed for source appendix generation: {exc}") from exc
        return SourceAppendixResult(
            status="skipped_malformed_ledger",
            warnings=["Source appendix skipped because claim_ledger.json was malformed."],
        )
    if result.source_count == 0:
        message = "No usable cited sources could be resolved for source appendix generation."
        if explicit:
            raise RuntimeError(message)
        result.status = "generated_with_warnings"
        result.warnings.append(message)
    if result.markdown:
        appendix_path.write_text(result.markdown, encoding="utf-8")
    return result


def _as_bool(value: Any, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "y", "on"}
    return bool(value)
