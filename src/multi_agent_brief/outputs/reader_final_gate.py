from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Literal


FindingKind = Literal[
    "src_marker",
    "bare_claim_id",
    "source_id",
    "process_wording",
    "blank_citation_row",
    "local_path",
    "debug_residue",
    "atom_id",
]

COUNT_KEYS = {
    "src_marker": "src_marker_count",
    "bare_claim_id": "bare_claim_id_count",
    "source_id": "source_id_count",
    "process_wording": "process_wording_count",
    "blank_citation_row": "blank_citation_row_count",
    "local_path": "local_path_count",
    "debug_residue": "debug_residue_count",
    "atom_id": "atom_id_count",
}

_SRC_MARKER_RE = re.compile(r"\[(?:src|source):[^\]]+\]", re.IGNORECASE)
_CLAIM_ID_RE = re.compile(
    r"(?<![A-Za-z0-9_])(?:\[(?:CLM-\d{3,}|CL-\d{3,})\]|CLM-\d{3,}|CL-\d{3,}|(?:[A-Z][A-Z0-9]*_)?CLAIM_[A-Z0-9][A-Z0-9_-]*)(?![A-Za-z0-9_])"
)
_SOURCE_ID_RE = re.compile(
    r"(?<![A-Za-z0-9_])(?:[A-Z][A-Z0-9]*_)?(?:SRC|SOURCE)_[A-Z0-9][A-Z0-9_-]*(?![A-Za-z0-9_])"
)
_CONTEXTUAL_SRC_ID_RE = re.compile(
    r"(?i)(?:source[_\s-]*id|source\s+ref(?:erence)?|来源\s*ID|源\s*ID)[:：\s`'\"]*(SRC-\d{3,})"
)
_LOCAL_PATH_RE = re.compile(r"(?:/Users/[^\s)]+|/mnt/data/[^\s)]+|file://[^\s)]+|[A-Za-z]:\\[^\s)]+)")
_DEBUG_RE = re.compile(r"\b(?:DEBUG|TRACE)\b")
_ATOM_ID_RE = re.compile(r"(?<![A-Za-z0-9_])AC-\d{4}-\d{2}(?![A-Za-z0-9_])")
_ATOM_ID_WORDING_RE = re.compile(r"\batom(?:\s+|-)+ids?\b", re.IGNORECASE)
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")

_PROCESS_WORDINGS = [
    "Analyst subagent",
    "Auditor subagent",
    "Claim Ledger",
    "source appendix generated from cited Claim Ledger",
    "Human review required before distribution",
    "audited_brief",
    "artifact_registry",
    "workflow_state",
    "quality_gate_report",
    "runtime_manifest",
    "agent_handoff",
    "claim_ledger.json",
    "finalize_report.json",
    "atomic_claim_graph",
    "Atomic Claim Graph",
    "atom_id",
    "事实账本",
    "声明账本",
    "分析师子代理",
    "审计师子代理",
    "审计员子代理",
    "运行交接单",
    "运行清单",
    "工作流状态",
    "产物注册表",
    "质量门禁",
]

_CITATION_SECTION_TITLES = [
    "citation index",
    "source index",
    "citation table",
    "references",
    "来源索引",
    "引用索引",
    "来源表",
]

_BLANK_CELL_VALUES = {"", "-", "--", "—", "n/a", "na", "null", "none", "无", "未提供", "unknown"}
_CITATION_ID_HEADERS = {
    "id",
    "claimid",
    "sourceid",
    "source",
    "sourceref",
    "sourcereference",
    "citationid",
    "citation",
    "referenceid",
    "reference",
    "refid",
    "ref",
    "编号",
    "来源",
    "来源编号",
    "来源id",
    "引用",
    "引用编号",
    "引用id",
    "证据编号",
}


@dataclass(frozen=True)
class ReaderResidueFinding:
    kind: FindingKind
    text: str
    line: int | None
    artifact: str
    message: str

    def to_dict(self) -> dict[str, object]:
        return asdict(self)


@dataclass(frozen=True)
class ReaderFinalGateResult:
    status: Literal["pass", "fail"]
    findings: list[ReaderResidueFinding]
    counts: dict[str, int]

    def to_report_dict(self) -> dict[str, object]:
        payload: dict[str, object] = {"status": self.status}
        payload.update(self.counts)
        payload["sample_findings"] = [
            finding.to_dict() for finding in self.findings[:10]
        ]
        return payload


def detect_reader_residue(
    markdown: str,
    artifact: str,
    *,
    allow_compliance_footer: bool = False,
) -> ReaderFinalGateResult:
    findings: list[ReaderResidueFinding] = []
    in_citation_section = False
    citation_table_header: list[str] | None = None
    source_table_header: list[str] | None = None

    for line_number, line in enumerate(markdown.splitlines(), start=1):
        heading = _HEADING_RE.match(line)
        if heading:
            title = heading.group(2).strip().lower()
            in_citation_section = any(marker in title for marker in _CITATION_SECTION_TITLES)
            citation_table_header = None
            source_table_header = None

        _collect_regex_findings(
            findings,
            kind="src_marker",
            regex=_SRC_MARKER_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains an internal source marker.",
        )
        _collect_regex_findings(
            findings,
            kind="bare_claim_id",
            regex=_CLAIM_ID_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains a raw internal claim ID.",
        )
        _collect_regex_findings(
            findings,
            kind="source_id",
            regex=_SOURCE_ID_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains a raw internal source ID.",
        )
        _collect_regex_findings(
            findings,
            kind="source_id",
            regex=_CONTEXTUAL_SRC_ID_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains an internal source ID.",
        )
        _collect_regex_findings(
            findings,
            kind="local_path",
            regex=_LOCAL_PATH_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains a local or file URL path.",
        )
        _collect_regex_findings(
            findings,
            kind="debug_residue",
            regex=_DEBUG_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains debug or trace residue.",
        )
        _collect_regex_findings(
            findings,
            kind="atom_id",
            regex=_ATOM_ID_RE,
            line=line,
            line_number=line_number,
            artifact=artifact,
            message="Reader-facing output contains an Atomic Claim Graph atom ID.",
        )
        _collect_process_wording_findings(
            findings,
            line=line,
            line_number=line_number,
            artifact=artifact,
            allow_compliance_footer=allow_compliance_footer,
        )
        cells = _table_cells(line)
        if cells is None:
            source_table_header = None
        elif not _TABLE_SEPARATOR_RE.match(line.strip()):
            if not in_citation_section:
                if source_table_header is None:
                    if _has_citation_id_header(cells):
                        source_table_header = cells
                elif _has_blank_citation_id_cell(source_table_header, cells):
                    findings.append(
                        ReaderResidueFinding(
                            kind="blank_citation_row",
                            text=_shorten(line.strip()),
                            line=line_number,
                            artifact=artifact,
                            message="Reader-facing table contains a blank source/reference cell.",
                        )
                    )
            if in_citation_section:
                if citation_table_header is None:
                    citation_table_header = cells
                elif _is_blank_citation_row_cells(cells):
                    findings.append(
                        ReaderResidueFinding(
                            kind="blank_citation_row",
                            text=_shorten(line.strip()),
                            line=line_number,
                            artifact=artifact,
                            message="Reader-facing source or citation section contains a blank table row.",
                        )
                    )
                elif _has_blank_citation_id_cell(citation_table_header, cells):
                    findings.append(
                        ReaderResidueFinding(
                            kind="blank_citation_row",
                            text=_shorten(line.strip()),
                            line=line_number,
                            artifact=artifact,
                            message="Reader-facing source or citation section contains a blank ID/source/reference cell.",
                        )
                    )

    counts = _empty_counts()
    for finding in findings:
        counts[COUNT_KEYS[finding.kind]] += 1
    return ReaderFinalGateResult(
        status="fail" if findings else "pass",
        findings=findings,
        counts=counts,
    )


def detect_reader_residue_in_docx(
    path: Path,
    *,
    artifact: str | None = None,
    allow_compliance_footer: bool = False,
) -> ReaderFinalGateResult:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return ReaderFinalGateResult(status="pass", findings=[], counts=_empty_counts())

    document = Document(str(path))
    text = "\n".join(_docx_text_parts(document))
    return detect_reader_residue(
        text,
        artifact=artifact or str(path),
        allow_compliance_footer=allow_compliance_footer,
    )


def combine_reader_final_gate_results(
    results: list[ReaderFinalGateResult],
) -> ReaderFinalGateResult:
    findings: list[ReaderResidueFinding] = []
    counts = _empty_counts()
    for result in results:
        findings.extend(result.findings)
        for key, value in result.counts.items():
            counts[key] = counts.get(key, 0) + value
    return ReaderFinalGateResult(
        status="fail" if findings else "pass",
        findings=findings,
        counts=counts,
    )


def _docx_text_parts(document: object) -> list[str]:
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    parts.extend(_table_text_parts(document.tables))
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            parts.extend(paragraph.text for paragraph in container.paragraphs)
            parts.extend(_table_text_parts(container.tables))
    return parts


def _table_text_parts(tables: object) -> list[str]:
    return [
        cell.text
        for table in tables
        for row in table.rows
        for cell in row.cells
    ]


def _collect_regex_findings(
    findings: list[ReaderResidueFinding],
    *,
    kind: FindingKind,
    regex: re.Pattern[str],
    line: str,
    line_number: int,
    artifact: str,
    message: str,
) -> None:
    for match in regex.finditer(line):
        text = match.group(0)
        findings.append(
            ReaderResidueFinding(
                kind=kind,
                text=_shorten(text),
                line=line_number,
                artifact=artifact,
                message=message,
            )
        )


def _collect_process_wording_findings(
    findings: list[ReaderResidueFinding],
    *,
    line: str,
    line_number: int,
    artifact: str,
    allow_compliance_footer: bool,
) -> None:
    line_lower = line.lower()
    for wording in _PROCESS_WORDINGS:
        if wording == "Human review required before distribution" and allow_compliance_footer:
            continue
        found = wording in line if _has_cjk(wording) else wording.lower() in line_lower
        if not found:
            continue
        findings.append(
            ReaderResidueFinding(
                kind="process_wording",
                text=_shorten(wording),
                line=line_number,
                artifact=artifact,
                message="Reader-facing output contains internal workflow/process wording.",
            )
        )
    for match in _ATOM_ID_WORDING_RE.finditer(line):
        findings.append(
            ReaderResidueFinding(
                kind="process_wording",
                text=_shorten(match.group(0)),
                line=line_number,
                artifact=artifact,
                message="Reader-facing output contains internal workflow/process wording.",
            )
        )


def _is_blank_citation_row(line: str) -> bool:
    cells = _table_cells(line)
    return bool(cells and _is_blank_citation_row_cells(cells))


def _table_cells(line: str) -> list[str] | None:
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return None
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    if len(cells) < 2:
        return None
    return cells


def _is_blank_citation_row_cells(cells: list[str]) -> bool:
    return all(cell.strip().lower() in _BLANK_CELL_VALUES for cell in cells)


def _has_blank_citation_id_cell(header_cells: list[str], row_cells: list[str]) -> bool:
    for index, header in enumerate(header_cells):
        if index >= len(row_cells):
            continue
        if _is_citation_id_header(header) and row_cells[index].strip().lower() in _BLANK_CELL_VALUES:
            return True
    return False


def _has_citation_id_header(header_cells: list[str]) -> bool:
    return any(_is_citation_id_header(header) for header in header_cells)


def _is_citation_id_header(header: str) -> bool:
    normalized = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", header.strip().lower())
    return normalized in _CITATION_ID_HEADERS


def _empty_counts() -> dict[str, int]:
    return {count_key: 0 for count_key in COUNT_KEYS.values()}


def _shorten(value: str, limit: int = 120) -> str:
    compact = " ".join(str(value).split())
    if len(compact) <= limit:
        return compact
    return compact[: limit - 1].rstrip() + "…"


def _has_cjk(value: str) -> bool:
    return any("\u4e00" <= char <= "\u9fff" for char in value)
